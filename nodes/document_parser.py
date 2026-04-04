import io
from dataclasses import dataclass, field
from typing import Literal
from nodes.base import BaseNode
import docx

@dataclass
class DocumentBlock:
    """
    A single translatable unit from the document.
    block_ref holds the live object (paragraph, cell) so the rebuilder
    can write directly back into it — no positional re-matching needed.
    """
    block_id: str
    block_type: Literal["paragraph", "heading", "table_cell", "text_frame"]
    source_text: str
    translated_text: str = ""
    block_ref: object = field(default=None, repr=False)  # live python-docx object
    metadata: dict = field(default_factory=dict)          # style name, level, etc.


class DocumentParserNode(BaseNode):

    async def run(self, context: dict) -> dict:
        filename: str = context.get("source_filename", "")
        file_bytes: bytes = context.get("file_bytes")  # raw bytes, added below

        if not file_bytes:
            # Fallback: already-extracted raw_text path (plain .txt)
            raw_text = context.get("raw_text", "")
            blocks = [DocumentBlock(
                block_id="b0",
                block_type="paragraph",
                source_text=raw_text,
            )]
            return {**context, "document_blocks": blocks, "segments": [raw_text]}

        if filename.endswith(".docx"):
            blocks = self._parse_docx(file_bytes)
        elif filename.endswith(".pdf"):
            blocks = self._parse_pdf(file_bytes)
        else:
            raw_text = context.get("raw_text", "")
            blocks = [DocumentBlock(
                block_id="b0", block_type="paragraph", source_text=raw_text
            )]

        # segments = all source texts, fed to RAG + LLM exactly as before
        segments = [b.source_text for b in blocks if b.source_text.strip()]

        return {
            **context,
            "document_blocks": blocks,
            "segments": segments,
            "segment_count": len(segments),
            # Keep raw_text for backwards compat with plain-text workflows
            "raw_text": "\n".join(segments),
        }

    # ── DOCX ─────────────────────────────────────────────────────────────────

    def _parse_docx(self, file_bytes: bytes) -> list[DocumentBlock]:
        doc = docx.Document(io.BytesIO(file_bytes))
        blocks = []
        idx = 0
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue  # empty or image-only paragraph — skip, don't create a block

            # Check this paragraph actually has translatable text runs
            # (not just drawing runs which also contribute to para.text weirdly)
            has_text_run = any(
                r.find(f"{{{W}}}t") is not None
                and r.find(f"{{{W}}}drawing") is None
                for r in para._element.findall(f"{{{W}}}r")
            )
            if not has_text_run:
                continue

            style_name = para.style.name if para.style else ""
            block_type = "heading" if style_name.startswith("Heading") else "paragraph"
            blocks.append(DocumentBlock(
                block_id=f"p{idx}",
                block_type=block_type,
                source_text=text,
                block_ref=para,
                metadata={"style": style_name},
            ))
            idx += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if not text:
                            continue
                        blocks.append(DocumentBlock(
                            block_id=f"tc{idx}",
                            block_type="table_cell",
                            source_text=text,
                            block_ref=para,
                            metadata={"style": para.style.name if para.style else ""},
                        ))
                        idx += 1

        if blocks:
            blocks[0].metadata["_docx_doc"] = doc

        return blocks

    # ── PDF ──────────────────────────────────────────────────────────────────

    def _parse_pdf(self, file_bytes: bytes) -> list[DocumentBlock]:
        """
        PDF strategy: extract text blocks with bbox metadata.
        Rebuilding is done with reportlab overlay on the original PDF.
        We store the original bytes for the rebuilder to use.
        """
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        blocks = []
        idx = 0

        for page_num, page in enumerate(reader.pages):
            # extract_text() with layout mode preserves reading order
            text = page.extract_text(extraction_mode="layout") or ""
            # Split into non-empty lines as individual blocks
            for line in text.splitlines():
                line = line.strip()
                if not line:
                    continue
                blocks.append(DocumentBlock(
                    block_id=f"pdf{idx}",
                    block_type="paragraph",
                    source_text=line,
                    block_ref=None,   # PDF blocks can't be mutated in place
                    metadata={"page": page_num, "_pdf_bytes": file_bytes if idx == 0 else None},
                ))
                idx += 1

        return blocks