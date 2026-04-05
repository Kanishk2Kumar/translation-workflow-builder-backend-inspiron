import hashlib
import hmac
import json
import uuid
from datetime import datetime
import base64
import os

from fastapi import APIRouter, BackgroundTasks, HTTPException, UploadFile, File, Form
from pydantic import BaseModel

from db import get_pool
from executor import execute_workflow
from nodes.document_parser import DocumentParserNode
from nodes.document_rebuilder import DocumentRebuilderNode
from nodes.output import OutputNode
from nodes.output import seed_translation_memory

router = APIRouter(prefix="/workflow", tags=["workflow"])
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}


class RunWorkflowResponse(BaseModel):
    execution_id: str
    status: str
    output: dict
    logs: list
    cache_hit: bool = False


class SegmentPair(BaseModel):
    segment_index: int
    source_text: str
    translated_text: str
    edited: bool = False


class ExecutionSegmentsResponse(BaseModel):
    execution_id: str
    workflow_id: str
    target_language: str
    segment_count: int
    segments: list[SegmentPair]


class SegmentEditIn(BaseModel):
    segment_index: int
    source_text: str
    translated_text: str


class RetranslateExecutionRequest(BaseModel):
    segments: list[SegmentEditIn]


def _normalize_workflow_row(row) -> dict:
    workflow = dict(row)
    workflow.pop("auth_token", None)
    for key in ("nodes", "edges"):
        value = workflow.get(key)
        if isinstance(value, str):
            try:
                workflow[key] = json.loads(value)
            except json.JSONDecodeError:
                pass
    return workflow


@router.get("/user/{user_id}")
async def list_user_workflows(user_id: str):
    pool = get_pool()
    rows = await pool.fetch(
        """
        SELECT id, agent_id, nodes, edges, created_at, updated_at, name, description, user_id, auth_type
        FROM workflows
        WHERE user_id = $1::uuid
        """,
        user_id,
    )
    return [_normalize_workflow_row(row) for row in rows]


def _normalize_auth_type(auth_type: str | None) -> str:
    return (auth_type or "none").strip().lower()


def _assert_workflow_auth_from_row(workflow_row, provided_auth_token: str | None) -> None:
    auth_type = _normalize_auth_type(workflow_row.get("auth_type") if workflow_row else None)
    if auth_type == "none":
        return

    if auth_type != "api_key":
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported workflow auth_type '{auth_type}'.",
        )

    expected_token = (workflow_row.get("auth_token") or "").strip()
    provided_token = (provided_auth_token or "").strip()

    if not provided_token:
        raise HTTPException(
            status_code=401,
            detail="This workflow requires an auth token.",
        )

    if not expected_token or not hmac.compare_digest(provided_token, expected_token):
        raise HTTPException(
            status_code=403,
            detail="Invalid workflow auth token.",
        )




# ─── Text extraction ──────────────────────────────────────────────────────────

def is_image_upload(filename: str, content_type: str | None = None) -> bool:
    _, ext = os.path.splitext(filename.lower())
    if ext in IMAGE_EXTENSIONS:
        return True
    return bool(content_type and content_type.startswith("image/"))


def extract_text(filename: str, contents: bytes, content_type: str | None = None) -> str:
    if is_image_upload(filename, content_type):
        return ""

    if filename.endswith(".txt"):
        return contents.decode("utf-8", errors="ignore")

    if filename.endswith(".pdf"):
        try:
            import io
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(contents))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            raise HTTPException(
                status_code=400,
                detail="PDF support requires pypdf. Run: pip install pypdf",
            )

    if filename.endswith(".docx"):
        try:
            import io
            import docx
            doc = docx.Document(io.BytesIO(contents))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise HTTPException(
                status_code=400,
                detail="DOCX support requires python-docx. Run: pip install python-docx",
            )

    raise HTTPException(
        status_code=400,
        detail=(
            f"Unsupported file type: {filename}. Supported: .txt, .pdf, .docx, "
            ".png, .jpg, .jpeg, .webp, .bmp, .tif, .tiff"
        ),
    )


# ─── Document hash helpers ────────────────────────────────────────────────────

def compute_hash(contents: bytes, target_language: str) -> str:
    """
    Hash = SHA256(file bytes + target language).
    Same file in different target languages = different cache entries.
    """
    h = hashlib.sha256()
    h.update(contents)
    h.update(target_language.encode())
    return h.hexdigest()


def _serialize_execution_input(
    *,
    filename: str,
    content_type: str | None,
    target_language: str,
    document_hash: str,
    contents: bytes,
    extra: dict | None = None,
) -> dict:
    payload = {
        "filename": filename,
        "content_type": content_type or "",
        "target_language": target_language,
        "document_hash": document_hash,
        "source_file_b64": base64.b64encode(contents).decode("utf-8"),
    }
    if extra:
        payload.update(extra)
    return payload


def _decode_execution_file_bytes(execution_input: dict) -> bytes:
    encoded = execution_input.get("source_file_b64")
    if not encoded:
        raise HTTPException(
            status_code=409,
            detail="This execution does not have the original source document stored, so it cannot be retranslated.",
        )

    try:
        return base64.b64decode(encoded)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="Stored source document could not be decoded for retranslation.",
        ) from exc


def _build_human_review_seed_payload(
    segments: list[SegmentEditIn],
    target_language: str,
    source_language: str = "en",
) -> dict:
    segment_translations: dict[str, str] = {}
    for segment in segments:
        if not segment.source_text.strip():
            continue
        segment_translations[segment.source_text] = segment.translated_text

    return {
        "source_language": source_language,
        "target_language": target_language,
        "segment_translations": segment_translations,
    }


async def get_cached_execution(
    pool, document_hash: str, workflow_id: str
) -> dict | None:
    """
    Look up a previous successful execution with the same document hash
    for the same workflow. Returns the output payload or None.
    """
    row = await pool.fetchrow(
        """
        SELECT id, output
        FROM executions
        WHERE workflow_id = $1
          AND status = 'success'
          AND input::jsonb->>'document_hash' = $2
        ORDER BY started_at DESC
        LIMIT 1
        """,
        workflow_id,
        document_hash,
    )
    if not row or not row["output"]:
        return None
    return {
        "execution_id": str(row["id"]),
        "output": json.loads(row["output"]) if isinstance(row["output"], str) else row["output"],
    }


async def _execute_workflow_run(
    *,
    pool,
    workflow_id: str,
    background_tasks: BackgroundTasks,
    contents: bytes,
    filename: str,
    content_type: str | None,
    target_language: str,
    auth_token: str | None = None,
    skip_cache: bool = False,
    execution_input_extra: dict | None = None,
) -> RunWorkflowResponse:
    document_hash = compute_hash(contents, target_language)

    raw_text = extract_text(filename, contents, content_type)

    row = await pool.fetchrow(
        "SELECT id, nodes, edges, auth_type, auth_token FROM workflows WHERE id = $1", workflow_id,
    )
    if not row:
        raise HTTPException(status_code=404, detail=f"Workflow '{workflow_id}' not found")
    _assert_workflow_auth_from_row(row, auth_token)

    if not skip_cache:
        cached = await get_cached_execution(pool, document_hash, workflow_id)
        if cached:
            return RunWorkflowResponse(
                execution_id=cached["execution_id"],
                status="success",
                output=cached["output"],
                logs=[],
                cache_hit=True,
            )

    nodes = json.loads(row["nodes"]) if isinstance(row["nodes"], str) else row["nodes"]
    edges = json.loads(row["edges"]) if isinstance(row["edges"], str) else row["edges"]

    if not nodes:
        raise HTTPException(status_code=400, detail="Workflow has no nodes")

    execution_id = str(uuid.uuid4())
    execution_input = _serialize_execution_input(
        filename=filename,
        content_type=content_type,
        target_language=target_language,
        document_hash=document_hash,
        contents=contents,
        extra=execution_input_extra,
    )

    await pool.execute(
        """
        INSERT INTO executions (id, workflow_id, status, input, started_at)
        VALUES ($1, $2, 'running', $3, $4)
        """,
        execution_id,
        workflow_id,
        json.dumps(execution_input),
        datetime.utcnow(),
    )

    initial_context = {
        "raw_text": raw_text,
        "original_raw_text": raw_text,
        "original_segments": [raw_text] if raw_text else [],
        "file_bytes": contents,
        "target_language": target_language,
        "execution_id": execution_id,
        "workflow_id": workflow_id,
        "source_filename": filename,
        "source_content_type": content_type or "",
        "document_hash": document_hash,
        "source_language": execution_input.get("source_language", "en"),
        "user_id": '37720c15-ff75-49eb-a538-b25fd2273d30',
    }

    try:
        final_context = await execute_workflow(
            nodes=nodes,
            edges=edges,
            initial_context=initial_context,
        )
    except Exception as e:
        await pool.execute(
            """
            UPDATE executions
            SET status = 'failed', logs = $1, completed_at = $2
            WHERE id = $3
            """,
            json.dumps([{"error": str(e)}]),
            datetime.utcnow(),
            execution_id,
        )
        raise HTTPException(status_code=500, detail=f"Workflow execution failed: {str(e)}")

    final_output = final_context.get("final_output", {})
    logs = final_context.get("_logs", [])
    tm_seed_payload = final_context.get("tm_seed_payload")

    if tm_seed_payload:
        background_tasks.add_task(seed_translation_memory, tm_seed_payload)

    output_doc_bytes = final_context.get("output_document_bytes")
    if output_doc_bytes:
        final_output["document_b64"] = base64.b64encode(output_doc_bytes).decode("utf-8")

    return RunWorkflowResponse(
        execution_id=execution_id,
        status="success",
        output=final_output,
        logs=logs,
        cache_hit=False,
    )


async def _load_execution_row(pool, workflow_id: str, execution_id: str):
    row = await pool.fetchrow(
        """
        SELECT input, output
        FROM executions
        WHERE id = $1 AND workflow_id = $2
        """,
        execution_id,
        workflow_id,
    )
    if not row:
        raise HTTPException(status_code=404, detail="Execution not found")
    return row


def _extract_segment_pairs_from_output(output: dict) -> list[dict]:
    segments = output.get("segments")
    if isinstance(segments, list):
        return segments

    translated_text = output.get("translated_text", "")
    if not translated_text:
        return []

    return [{
        "segment_index": 1,
        "source_text": "",
        "translated_text": translated_text,
        "edited": False,
    }]


def _apply_segment_edits(existing_segments: list[dict], edits: list[SegmentEditIn]) -> tuple[list[dict], int]:
    updated_segments = [dict(segment) for segment in existing_segments]
    edited_count = 0

    for edit in edits:
        idx = edit.segment_index - 1
        if idx < 0 or idx >= len(updated_segments):
            raise HTTPException(
                status_code=400,
                detail=f"segment_index {edit.segment_index} is out of range for this execution.",
            )

        current_segment = updated_segments[idx]
        expected_source = current_segment.get("source_text", "")
        if edit.source_text and expected_source and edit.source_text != expected_source:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"source_text mismatch for segment_index {edit.segment_index}. "
                    "The segment no longer matches the stored execution output."
                ),
            )

        current_segment["translated_text"] = edit.translated_text
        current_segment["edited"] = True
        edited_count += 1

    return updated_segments, edited_count


async def _create_manual_retranslation_execution(
    *,
    pool,
    workflow_id: str,
    parent_execution_id: str,
    execution_input: dict,
    parent_output: dict,
    updated_segments: list[dict],
) -> RunWorkflowResponse:
    filename = execution_input.get("filename")
    if not filename:
        raise HTTPException(
            status_code=409,
            detail="Original filename is missing from this execution, so it cannot be rebuilt.",
        )

    source_file_bytes = _decode_execution_file_bytes(execution_input)
    target_language = execution_input.get("target_language", parent_output.get("target_language", "hi"))
    source_language = execution_input.get("source_language", parent_output.get("source_language", "en"))
    document_hash = compute_hash(source_file_bytes, target_language)
    execution_id = str(uuid.uuid4())

    execution_input_payload = _serialize_execution_input(
        filename=filename,
        content_type=execution_input.get("content_type"),
        target_language=target_language,
        document_hash=document_hash,
        contents=source_file_bytes,
        extra={
            "source_language": source_language,
            "parent_execution_id": parent_execution_id,
            "human_reviewed": True,
            "human_review_segment_count": sum(1 for segment in updated_segments if segment.get("edited")),
            "retranslation_mode": "manual_patch",
        },
    )

    await pool.execute(
        """
        INSERT INTO executions (id, workflow_id, status, input, started_at)
        VALUES ($1, $2, 'running', $3, $4)
        """,
        execution_id,
        workflow_id,
        json.dumps(execution_input_payload),
        datetime.utcnow(),
    )

    parser = DocumentParserNode(node_id="manual_document_parser", config={})
    parsed_context = await parser.run({
        "source_filename": filename,
        "file_bytes": source_file_bytes,
        "raw_text": "",
    })

    document_blocks = parsed_context.get("document_blocks", [])
    if len(document_blocks) != len(updated_segments):
        raise HTTPException(
            status_code=409,
            detail=(
                "Stored execution segments no longer match the parsed source document. "
                "A full rerun would be required for this file."
            ),
        )

    for block, segment in zip(document_blocks, updated_segments):
        block.translated_text = segment.get("translated_text", "")

    rebuilder = DocumentRebuilderNode(node_id="manual_document_rebuilder", config={})
    if filename.endswith(".docx"):
        output_document_bytes = rebuilder._rebuild_docx(document_blocks)
        output_document_format = "docx"
    elif filename.endswith(".pdf"):
        output_document_bytes = rebuilder._rebuild_pdf(document_blocks, parsed_context)
        output_document_format = "pdf"
    else:
        raise HTTPException(
            status_code=400,
            detail="Manual retranslation rebuild currently supports DOCX and PDF source files only.",
        )

    translated_text = "\n".join(segment.get("translated_text", "") for segment in updated_segments)
    logs = [{
        "node_id": "manual_review_patch",
        "node_type": "manual_review_patch",
        "status": "success",
        "edited_segment_count": sum(1 for segment in updated_segments if segment.get("edited")),
        "parent_execution_id": parent_execution_id,
    }]

    output_node = OutputNode(node_id="manual_output", config={"include_audit": True})
    final_context = await output_node.run({
        "_logs": logs,
        "execution_id": execution_id,
        "workflow_id": workflow_id,
        "target_language": target_language,
        "source_language": source_language,
        "source_filename": filename,
        "segment_count": len(updated_segments),
        "translated_text": translated_text,
        "llm_model": "human_review_patch",
        "input_tokens": 0,
        "output_tokens": 0,
        "tm_hit": False,
        "output_document_bytes": output_document_bytes,
        "output_document_format": output_document_format,
        "rag_stats": parent_output.get("rag_stats", {}),
        "review_required": False,
        "review_reason": "Human segment edits applied without rerunning the workflow.",
        "compliance_status": parent_output.get("compliance_status"),
        "compliance_errors": parent_output.get("compliance_errors", []),
        "compliance_suggestions": parent_output.get("compliance_suggestions", []),
        "compliance_report": parent_output.get("compliance_report"),
        "ocr_provider": parent_output.get("ocr_provider"),
        "ocr_confidence": parent_output.get("ocr_confidence"),
        "ocr_status": parent_output.get("ocr_status"),
        "ocr_warnings": parent_output.get("ocr_warnings", []),
        "original_segments": [segment.get("source_text", "") for segment in updated_segments],
        "segment_translations": {},
        "output_segments_override": updated_segments,
    })

    final_output = final_context.get("final_output", {})
    final_output["segments"] = updated_segments
    final_output["translated_text"] = translated_text
    if output_document_bytes:
        final_output["document_b64"] = base64.b64encode(output_document_bytes).decode("utf-8")

    await pool.execute(
        """
        UPDATE executions
        SET output = $1
        WHERE id = $2
        """,
        json.dumps({k: v for k, v in final_output.items() if k != "document_b64"}),
        execution_id,
    )

    return RunWorkflowResponse(
        execution_id=execution_id,
        status="success",
        output=final_output,
        logs=logs,
        cache_hit=False,
    )


# ─── POST /workflow/{workflow_id}/run ─────────────────────────────────────────

@router.post("/{workflow_id}/run", response_model=RunWorkflowResponse)
async def run_workflow(
    workflow_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    target_language: str = Form(default="hi"),
    auth_token: str | None = Form(default=None),
):
    pool = get_pool()
    contents = await file.read()

    return await _execute_workflow_run(
        pool=pool,
        workflow_id=workflow_id,
        background_tasks=background_tasks,
        contents=contents,
        filename=file.filename,
        content_type=file.content_type,
        target_language=target_language,
        auth_token=auth_token,
    )

# ─── GET /execution/{execution_id}/status ────────────────────────────────────

@router.get("/{workflow_id}/execution/{execution_id}/segments", response_model=ExecutionSegmentsResponse)
async def get_execution_segments(
    workflow_id: str,
    execution_id: str,
):
    pool = get_pool()
    row = await _load_execution_row(pool, workflow_id, execution_id)

    execution_input = json.loads(row["input"]) if isinstance(row["input"], str) else row["input"] or {}
    output = json.loads(row["output"]) if isinstance(row["output"], str) else row["output"] or {}
    segments = _extract_segment_pairs_from_output(output)

    return ExecutionSegmentsResponse(
        execution_id=execution_id,
        workflow_id=workflow_id,
        target_language=output.get("target_language") or execution_input.get("target_language", "hi"),
        segment_count=len(segments),
        segments=[SegmentPair(**segment) for segment in segments],
    )


@router.post("/{workflow_id}/execution/{execution_id}/retranslate", response_model=RunWorkflowResponse)
async def retranslate_workflow(
    workflow_id: str,
    execution_id: str,
    request: RetranslateExecutionRequest,
    background_tasks: BackgroundTasks,
):
    pool = get_pool()
    if not request.segments:
        raise HTTPException(status_code=400, detail="At least one edited segment is required.")

    row = await _load_execution_row(pool, workflow_id, execution_id)
    execution_input = json.loads(row["input"]) if isinstance(row["input"], str) else row["input"] or {}
    parent_output = json.loads(row["output"]) if isinstance(row["output"], str) else row["output"] or {}
    existing_segments = _extract_segment_pairs_from_output(parent_output)
    if not existing_segments:
        raise HTTPException(
            status_code=409,
            detail="This execution does not contain segment output, so manual retranslation cannot patch it.",
        )

    updated_segments, edited_count = _apply_segment_edits(existing_segments, request.segments)
    reviewed_tm_payload = _build_human_review_seed_payload(
        segments=request.segments,
        target_language=execution_input.get("target_language", "hi"),
        source_language=execution_input.get("source_language", "en"),
    )
    if reviewed_tm_payload["segment_translations"]:
        background_tasks.add_task(seed_translation_memory, reviewed_tm_payload)

    response = await _create_manual_retranslation_execution(
        pool=pool,
        workflow_id=workflow_id,
        parent_execution_id=execution_id,
        execution_input=execution_input,
        parent_output=parent_output,
        updated_segments=updated_segments,
    )
    response.logs[0]["edited_segment_count"] = edited_count
    return response



@router.get("/{workflow_id}/execution/{execution_id}/download")
async def download_translated_document(
    workflow_id: str,
    execution_id: str,
):
    import re
    from fastapi.responses import Response

    pool = get_pool()
    row = await pool.fetchrow(
        "SELECT output FROM executions WHERE id = $1 AND workflow_id = $2",
        execution_id, workflow_id,
    )
    if not row or not row["output"]:
        raise HTTPException(status_code=404, detail="Execution not found")

    output = json.loads(row["output"]) if isinstance(row["output"], str) else row["output"]

    doc_path: str | None = output.get("document_path")
    doc_format: str | None = output.get("document_format")
    phi_map: dict = output.get("phi_map", {})

    if not doc_path or not doc_format:
        raise HTTPException(
            status_code=404,
            detail="No document output for this execution — was a DOCX or PDF uploaded?"
        )

    if not os.path.exists(doc_path):
        raise HTTPException(
            status_code=410,
            detail="Translated document has expired from temporary storage"
        )

    with open(doc_path, "rb") as f:
        file_bytes = f.read()

    # ── Safety net: re-apply PHI restore directly to document bytes ──────────
    if phi_map:
        if doc_format == "docx":
            file_bytes = _restore_phi_in_docx(file_bytes, phi_map)
        # PDF restore can be added later

    content_types = {
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }

    return Response(
        content=file_bytes,
        media_type=content_types.get(doc_format, "application/octet-stream"),
        headers={
            "Content-Disposition": f"attachment; filename=translated.{doc_format}",
            "Content-Length": str(len(file_bytes)),
        },
    )


def _restore_phi_in_docx(file_bytes: bytes, phi_map: dict) -> bytes:
    import io
    import docx

    if not phi_map:
        return file_bytes

    doc = docx.Document(io.BytesIO(file_bytes))
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def restore_text(text: str) -> str:
        for placeholder, original in phi_map.items():
            text = text.replace(placeholder, original)
        return text

    for para in doc.paragraphs:
        for run in para.runs:
            if run.text and "PHIMASK" in run.text:
                run.text = restore_text(run.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text and "PHIMASK" in run.text:
                            run.text = restore_text(run.text)

    # Direct XML walk catches placeholders split across <w:t> elements
    for elem in doc.element.iter(f"{{{W}}}t"):
        if elem.text and "PHIMASK" in elem.text:
            elem.text = restore_text(elem.text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
